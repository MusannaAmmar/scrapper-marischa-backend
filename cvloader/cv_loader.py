from pypdf import PdfReader
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import zipfile
import os
from lxml import etree
import re



class FileLoader:
    def __init__(self, file_path):
        self.text = self.load(file_path)

    def load(self, file_path):
        """Load text from PDF or DOCX file based on file extension."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            return self.load_pdf(file_path)
        elif ext == '.docx':
            return self.load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: '{ext}'. Only .pdf and .docx are supported.")

    def load_pdf(self, pdf_file_path):
        """Extract text from PDF file given its path."""
        pdf_reader = PdfReader(pdf_file_path)
        text = ''
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
        return text

    def load_docx(self, docx_file_path):
        """Extract text from DOCX file, handling corrupted images gracefully."""
        try:
            # Try normal loading first
            doc = Document(docx_file_path)
            return self._extract_text(doc)

        except zipfile.BadZipFile:
            # File has corrupted media (images) — extract XML directly and parse text
            print("⚠ Corrupted media found in DOCX, extracting text only...")
            return self._extract_text_from_corrupt_docx(docx_file_path)

        except PackageNotFoundError as e:
            raise ValueError(f"Could not open DOCX file: {e}")

    def _extract_text(self, doc):
        """Extract text from a loaded Document object."""
        text = ''
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + '\n'
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + '\n'
        return text

    def _extract_text_from_corrupt_docx(self, docx_file_path):
        """Extract text from a DOCX with corrupted media by reading XML directly."""
      

        text = ''
        try:
            with zipfile.ZipFile(docx_file_path, 'r') as z:
                # Only read word/document.xml — skip media files
                if 'word/document.xml' in z.namelist():
                    with z.open('word/document.xml') as xml_file:
                        tree = etree.parse(xml_file)
                        # Extract all text nodes
                        root = tree.getroot()
                        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        paragraphs = root.findall('.//w:p', ns)
                        for para in paragraphs:
                            runs = para.findall('.//w:t', ns)
                            para_text = ''.join(r.text or '' for r in runs)
                            if para_text.strip():
                                text += para_text + '\n'
        except Exception as e:
            raise ValueError(f"Failed to extract text from corrupted DOCX: {e}")

        return text



# loader = FileLoader(r"C:\Users\buzz\Downloads\CV Engineer Abdullatif.docx")
# loader= FileLoader(r"C:\Users\buzz\Desktop\C.V. Marischa van Zantvoort.pdf")
# print(loader.text)